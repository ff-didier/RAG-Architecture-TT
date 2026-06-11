from pathlib import Path

import pytest
from docx import Document as DocxDocument

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture
def fixtures_dir() -> Path:
    FIXTURES_DIR.mkdir(exist_ok=True)
    return FIXTURES_DIR


@pytest.fixture
def sample_docx(fixtures_dir: Path) -> Path:
    path = fixtures_dir / "sample.docx"
    if not path.exists():
        doc = DocxDocument()
        doc.add_paragraph("RAG systems combine retrieval with generation.")
        doc.add_paragraph("Hybrid search merges BM25 and vector similarity.")
        doc.save(path)
    return path


@pytest.fixture
def sample_pdf(fixtures_dir: Path) -> Path:
    path = fixtures_dir / "sample.pdf"
    if not path.exists():
        # Minimal PDF with extractable text (valid enough for pypdf)
        pdf_bytes = b"""%PDF-1.4
1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj
2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj
3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200]
/Contents 4 0 R /Resources<< /Font<< /F1 5 0 R >> >> >>endobj
4 0 obj<< /Length 68 >>stream
BT /F1 12 Tf 10 100 Td (PDF content about pgvector and embeddings.) Tj ET
endstream
endobj
5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000386 00000 n 
trailer<< /Size 6 /Root 1 0 R >>
startxref
464
%%EOF"""
        path.write_bytes(pdf_bytes)
    return path
